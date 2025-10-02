#!/usr/bin/env python3
"""
Sub2Trans - Convert subtitle files to well-formatted transcripts

Copyright (c) 2024, Sub2Trans Contributors
All rights reserved.

Redistribution and use in source and binary forms, with or without
modification, are permitted provided that the following conditions are met:

1. Redistributions of source code must retain the above copyright notice, this
   list of conditions and the following disclaimer.

2. Redistributions in binary form must reproduce the above copyright notice,
   this list of conditions and the following disclaimer in the documentation
   and/or other materials provided with the distribution.

3. Neither the name of the copyright holder nor the names of its
   contributors may be used to endorse or promote products derived from
   this software without specific prior written permission.

THIS SOFTWARE IS PROVIDED BY THE COPYRIGHT HOLDERS AND CONTRIBUTORS "AS IS"
AND ANY EXPRESS OR IMPLIED WARRANTIES, INCLUDING, BUT NOT LIMITED TO, THE
IMPLIED WARRANTIES OF MERCHANTABILITY AND FITNESS FOR A PARTICULAR PURPOSE ARE
DISCLAIMED. IN NO EVENT SHALL THE COPYRIGHT HOLDER OR CONTRIBUTORS BE LIABLE
FOR ANY DIRECT, INDIRECT, INCIDENTAL, SPECIAL, EXEMPLARY, OR CONSEQUENTIAL
DAMAGES (INCLUDING, BUT NOT LIMITED TO, PROCUREMENT OF SUBSTITUTE GOODS OR
SERVICES; LOSS OF USE, DATA, OR PROFITS; OR BUSINESS INTERRUPTION) HOWEVER
CAUSED AND ON ANY THEORY OF LIABILITY, WHETHER IN CONTRACT, STRICT LIABILITY,
OR TORT (INCLUDING NEGLIGENCE OR OTHERWISE) ARISING IN ANY WAY OUT OF THE USE
OF THIS SOFTWARE, EVEN IF ADVISED OF THE POSSIBILITY OF SUCH DAMAGE.

Converts subtitle files to well-formatted transcripts with:
- Multi-format subtitle support (SRT, VTT, ASS, SSA, SBV, TTML, SAMI, LRC)
- Multi-platform video integration (YouTube, Vimeo, Twitch, etc.)
- AI-enhanced paragraph grouping and headline generation
- Multiple output formats (Markdown, HTML, PDF, DOCX, RTF)
- Configuration system with multiple AI provider support
"""

import re
import argparse
import sys
import json
import requests
import time
import subprocess
import tempfile
import os
from datetime import datetime, timedelta
from typing import List, Dict, Tuple, Optional
from dataclasses import dataclass
from tqdm import tqdm
from urllib.parse import urlparse, parse_qs
from pathlib import Path

# Global verbosity setting
VERBOSE = False


class ConfigManager:
    """Manages configuration file loading and AI provider settings"""
    
    def __init__(self, config_path: Optional[str] = None):
        self.config_path = config_path or self._get_default_config_path()
        self.config = self._load_config()
    
    def _get_default_config_path(self) -> str:
        """Get the default configuration file path in user's home directory"""
        home_dir = Path.home()
        config_file = home_dir / "sub2trans_config.json"
        return str(config_file)
    
    def _load_config(self) -> Dict:
        """Load configuration from file or create default"""
        if os.path.exists(self.config_path):
            try:
                with open(self.config_path, 'r', encoding='utf-8') as f:
                    return json.load(f)
            except Exception as e:
                log(f"Error loading config file: {e}", "ERROR")
                return self._get_default_config()
        else:
            # Create default config file
            default_config = self._get_default_config()
            try:
                os.makedirs(os.path.dirname(self.config_path), exist_ok=True)
                with open(self.config_path, 'w', encoding='utf-8') as f:
                    json.dump(default_config, f, indent=2)
                log(f"Created default config file: {self.config_path}")
            except Exception as e:
                log(f"Error creating config file: {e}", "ERROR")
            return default_config
    
    def _get_default_config(self) -> Dict:
        """Get default configuration"""
        return {
            "ai_providers": {
                "lm_studio": {
                    "enabled": True,
                    "base_url": "http://localhost:1234",
                    "default_model": "qwen/qwen3-4b-2507",
                    "timeout": 120,
                    "max_retries": 3
                },
                "openai": {
                    "enabled": False,
                    "api_key": "",
                    "base_url": "https://api.openai.com/v1",
                    "default_model": "gpt-3.5-turbo",
                    "timeout": 60,
                    "max_retries": 3
                },
                "anthropic": {
                    "enabled": False,
                    "api_key": "",
                    "base_url": "https://api.anthropic.com",
                    "default_model": "claude-3-haiku-20240307",
                    "timeout": 60,
                    "max_retries": 3
                },
                "grok": {
                    "enabled": False,
                    "api_key": "",
                    "base_url": "https://api.x.ai/v1",
                    "default_model": "grok-beta",
                    "timeout": 60,
                    "max_retries": 3
                },
                "google": {
                    "enabled": False,
                    "api_key": "",
                    "base_url": "https://generativelanguage.googleapis.com/v1beta",
                    "default_model": "gemini-pro",
                    "timeout": 60,
                    "max_retries": 3
                }
            },
            "default_provider": "lm_studio",
            "processing": {
                "max_gap_seconds": 3.0,
                "min_paragraph_length": 10,
                "wait_for_model": False,
                "no_headlines": False,
                "verbose": False
            },
            "output": {
                "default_format": "markdown",
                "pandoc_path": "pandoc",
                "auto_detect_format": True,
                "preferred_formats": ["markdown", "html", "pdf", "docx"]
            }
        }
    
    def get_provider_config(self, provider: str) -> Optional[Dict]:
        """Get configuration for a specific AI provider"""
        return self.config.get("ai_providers", {}).get(provider)
    
    def get_enabled_providers(self) -> List[str]:
        """Get list of enabled AI providers"""
        providers = []
        for provider, config in self.config.get("ai_providers", {}).items():
            if config.get("enabled", False):
                providers.append(provider)
        return providers
    
    def get_default_provider(self) -> str:
        """Get the default AI provider"""
        return self.config.get("default_provider", "lm_studio")
    
    def get_processing_config(self) -> Dict:
        """Get processing configuration"""
        return self.config.get("processing", {})
    
    def get_output_config(self) -> Dict:
        """Get output configuration"""
        return self.config.get("output", {})
    
    def update_provider_config(self, provider: str, updates: Dict) -> None:
        """Update configuration for a specific provider"""
        if "ai_providers" not in self.config:
            self.config["ai_providers"] = {}
        if provider not in self.config["ai_providers"]:
            self.config["ai_providers"][provider] = {}
        
        self.config["ai_providers"][provider].update(updates)
        self._save_config()
    
    def _save_config(self) -> None:
        """Save configuration to file"""
        try:
            with open(self.config_path, 'w', encoding='utf-8') as f:
                json.dump(self.config, f, indent=2)
        except Exception as e:
            log(f"Error saving config file: {e}", "ERROR")

def log(message: str, level: str = "INFO"):
    """Log message based on verbosity level"""
    if VERBOSE or level == "ERROR":
        print(message)

def log_progress(message: str):
    """Log progress message (always shown)"""
    print(message)

def handle_error(error: Exception, context: str = "", fallback_action: str = None) -> bool:
    """Handle errors gracefully with fallback options"""
    error_msg = f"Error in {context}: {str(error)}" if context else str(error)
    log(error_msg, "ERROR")
    
    if fallback_action:
        log(f"Attempting fallback: {fallback_action}")
        return True
    return False


class VideoURLHandler:
    """Handles video URL processing and timestamp generation"""
    
    def __init__(self, video_url: str = None):
        self.video_url = video_url
        self.platform = self._detect_platform()
        self.base_url = self._get_base_url()
    
    def _detect_platform(self) -> Optional[str]:
        """Detect video platform from URL"""
        if not self.video_url:
            return None
        
        url = self.video_url.lower()
        
        if 'youtube.com' in url or 'youtu.be' in url:
            return 'youtube'
        elif 'vimeo.com' in url:
            return 'vimeo'
        elif 'twitch.tv' in url or 'twitch.com' in url:
            return 'twitch'
        elif 'dailymotion.com' in url or 'dai.ly' in url:
            return 'dailymotion'
        elif 'tiktok.com' in url:
            return 'tiktok'
        elif 'instagram.com' in url or 'instagr.am' in url:
            return 'instagram'
        elif 'facebook.com' in url or 'fb.watch' in url:
            return 'facebook'
        elif 'twitter.com' in url or 'x.com' in url or 't.co' in url:
            return 'twitter'
        elif 'linkedin.com' in url:
            return 'linkedin'
        elif 'rumble.com' in url:
            return 'rumble'
        elif 'odysee.com' in url:
            return 'odysee'
        else:
            return 'unknown'
    
    def _get_base_url(self) -> Optional[str]:
        """Get the base video URL without existing timestamps"""
        if not self.video_url or self.platform == 'unknown':
            return None
        
        if self.platform == 'youtube':
            # Clean YouTube URL
            if 'youtu.be' in self.video_url:
                video_id = self.video_url.split('/')[-1].split('?')[0]
                return f"https://www.youtube.com/watch?v={video_id}"
            else:
                # Remove existing timestamp parameters
                parsed = urlparse(self.video_url)
                if parsed.query:
                    params = parse_qs(parsed.query)
                    if 'v' in params:
                        return f"https://www.youtube.com/watch?v={params['v'][0]}"
                return self.video_url
        
        elif self.platform == 'vimeo':
            # Clean Vimeo URL
            parsed = urlparse(self.video_url)
            return f"https://vimeo.com{parsed.path}"
        
        elif self.platform == 'twitch':
            # Clean Twitch URL - remove timestamp parameters
            parsed = urlparse(self.video_url)
            if parsed.query:
                params = parse_qs(parsed.query)
                # Remove timestamp parameters
                clean_params = {k: v for k, v in params.items() if k not in ['t', 'time']}
                if clean_params:
                    query_string = '&'.join([f"{k}={v[0]}" for k, v in clean_params.items()])
                    return f"{parsed.scheme}://{parsed.netloc}{parsed.path}?{query_string}"
            return f"{parsed.scheme}://{parsed.netloc}{parsed.path}"
        
        elif self.platform == 'dailymotion':
            # Clean Dailymotion URL
            parsed = urlparse(self.video_url)
            return f"{parsed.scheme}://{parsed.netloc}{parsed.path}"
        
        elif self.platform in ['tiktok', 'instagram', 'facebook', 'twitter', 'linkedin']:
            # For social media platforms, return the URL as-is (timestamp support varies)
            return self.video_url
        
        elif self.platform == 'rumble':
            # Clean Rumble URL
            parsed = urlparse(self.video_url)
            return f"{parsed.scheme}://{parsed.netloc}{parsed.path}"
        
        elif self.platform == 'odysee':
            # Clean Odysee URL
            parsed = urlparse(self.video_url)
            return f"{parsed.scheme}://{parsed.netloc}{parsed.path}"
        
        return self.video_url
    
    def get_timestamp_url(self, seconds: float) -> Optional[str]:
        """Generate timestamped URL for the given seconds"""
        if not self.base_url:
            return None
        
        if self.platform == 'youtube':
            return f"{self.base_url}&t={int(seconds)}s"
        elif self.platform == 'vimeo':
            return f"{self.base_url}#t={seconds:.3f}"
        elif self.platform == 'twitch':
            # Twitch uses ?t= parameter for timestamps
            return f"{self.base_url}?t={int(seconds)}s"
        elif self.platform == 'dailymotion':
            # Dailymotion uses ?start= parameter for timestamps
            return f"{self.base_url}?start={int(seconds)}"
        elif self.platform == 'tiktok':
            # TikTok doesn't support direct timestamp URLs, return base URL
            return self.base_url
        elif self.platform == 'instagram':
            # Instagram doesn't support direct timestamp URLs, return base URL
            return self.base_url
        elif self.platform == 'facebook':
            # Facebook uses ?t= parameter for timestamps
            return f"{self.base_url}?t={int(seconds)}"
        elif self.platform == 'twitter':
            # Twitter/X doesn't support direct timestamp URLs, return base URL
            return self.base_url
        elif self.platform == 'linkedin':
            # LinkedIn doesn't support direct timestamp URLs, return base URL
            return self.base_url
        elif self.platform == 'rumble':
            # Rumble uses ?t= parameter for timestamps
            return f"{self.base_url}?t={int(seconds)}"
        elif self.platform == 'odysee':
            # Odysee uses ?t= parameter for timestamps
            return f"{self.base_url}?t={int(seconds)}"
        else:
            return None
    
    def is_valid(self) -> bool:
        """Check if the video URL is valid and supported"""
        return self.platform in [
            'youtube', 'vimeo', 'twitch', 'dailymotion', 'tiktok', 
            'instagram', 'facebook', 'twitter', 'linkedin', 'rumble', 'odysee'
        ]


class AIClient:
    """Unified client for communicating with various AI providers"""
    
    def __init__(self, provider: str, config: Dict):
        self.provider = provider
        self.config = config
        self.base_url = config.get("base_url", "")
        self.model = config.get("default_model", "")
        self.api_key = config.get("api_key", "")
        self.timeout = config.get("timeout", 60)
        self.max_retries = config.get("max_retries", 3)
        
        self.session = requests.Session()
        self._setup_headers()
    
    def _setup_headers(self):
        """Setup headers based on provider"""
        if self.provider == "openai":
            self.session.headers.update({
                'Content-Type': 'application/json',
                'Authorization': f'Bearer {self.api_key}'
            })
        elif self.provider == "anthropic":
            self.session.headers.update({
                'Content-Type': 'application/json',
                'x-api-key': self.api_key,
                'anthropic-version': '2023-06-01'
            })
        elif self.provider == "grok":
            self.session.headers.update({
                'Content-Type': 'application/json',
                'Authorization': f'Bearer {self.api_key}'
            })
        elif self.provider == "google":
            self.session.headers.update({
                'Content-Type': 'application/json'
            })
        else:  # LM Studio and others
            self.session.headers.update({
                'Content-Type': 'application/json',
                'Accept': 'application/json'
            })
    
    def generate_text(self, prompt: str, max_tokens: int = 1000, temperature: float = 0.7, retries: int = None) -> str:
        """Generate text using various AI providers with retry mechanism"""
        if retries is None:
            retries = self.max_retries
            
        for attempt in range(retries):
            try:
                if self.provider == "openai":
                    return self._call_openai(prompt, max_tokens, temperature)
                elif self.provider == "anthropic":
                    return self._call_anthropic(prompt, max_tokens, temperature)
                elif self.provider == "grok":
                    return self._call_grok(prompt, max_tokens, temperature)
                elif self.provider == "google":
                    return self._call_google(prompt, max_tokens, temperature)
                else:  # LM Studio
                    return self._call_lm_studio(prompt, max_tokens, temperature)
                    
            except requests.exceptions.Timeout:
                log(f"Request timeout (attempt {attempt + 1}/{retries})")
            except requests.exceptions.ConnectionError:
                log(f"Connection error (attempt {attempt + 1}/{retries})")
            except Exception as e:
                log(f"Unexpected error: {e}")
            
            if attempt < retries - 1:
                log(f"Retrying in 2 seconds...")
                time.sleep(2)
        
        log("All retry attempts failed")
        return ""
    
    def _call_openai(self, prompt: str, max_tokens: int, temperature: float) -> str:
        """Call OpenAI API"""
        payload = {
            "model": self.model,
            "messages": [{"role": "user", "content": prompt}],
            "max_tokens": max_tokens,
            "temperature": temperature
        }
        
        response = self.session.post(
            f"{self.base_url}/chat/completions",
            json=payload,
            timeout=self.timeout
        )
        
        if response.status_code == 200:
            result = response.json()
            return result.get('choices', [{}])[0].get('message', {}).get('content', '').strip()
        else:
            raise Exception(f"OpenAI API error: {response.status_code} - {response.text}")
    
    def _call_anthropic(self, prompt: str, max_tokens: int, temperature: float) -> str:
        """Call Anthropic Claude API"""
        payload = {
            "model": self.model,
            "max_tokens": max_tokens,
            "temperature": temperature,
            "messages": [{"role": "user", "content": prompt}]
        }
        
        response = self.session.post(
            f"{self.base_url}/messages",
            json=payload,
            timeout=self.timeout
        )
        
        if response.status_code == 200:
            result = response.json()
            return result.get('content', [{}])[0].get('text', '').strip()
        else:
            raise Exception(f"Anthropic API error: {response.status_code} - {response.text}")
    
    def _call_grok(self, prompt: str, max_tokens: int, temperature: float) -> str:
        """Call Grok API"""
        payload = {
            "model": self.model,
            "messages": [{"role": "user", "content": prompt}],
            "max_tokens": max_tokens,
            "temperature": temperature
        }
        
        response = self.session.post(
            f"{self.base_url}/chat/completions",
            json=payload,
            timeout=self.timeout
        )
        
        if response.status_code == 200:
            result = response.json()
            return result.get('choices', [{}])[0].get('message', {}).get('content', '').strip()
        else:
            raise Exception(f"Grok API error: {response.status_code} - {response.text}")
    
    def _call_google(self, prompt: str, max_tokens: int, temperature: float) -> str:
        """Call Google Gemini API"""
        payload = {
            "contents": [{"parts": [{"text": prompt}]}],
            "generationConfig": {
                "maxOutputTokens": max_tokens,
                "temperature": temperature
            }
        }
        
        response = self.session.post(
            f"{self.base_url}/models/{self.model}:generateContent?key={self.api_key}",
            json=payload,
            timeout=self.timeout
        )
        
        if response.status_code == 200:
            result = response.json()
            return result.get('candidates', [{}])[0].get('content', {}).get('parts', [{}])[0].get('text', '').strip()
        else:
            raise Exception(f"Google API error: {response.status_code} - {response.text}")
    
    def _call_lm_studio(self, prompt: str, max_tokens: int, temperature: float) -> str:
        """Call LM Studio API (existing implementation)"""
        payload = {
            "model": self.model,
            "messages": [{"role": "user", "content": prompt}],
            "max_tokens": max_tokens,
            "temperature": temperature,
            "stream": False
        }
        
        response = self.session.post(
            f"{self.base_url}/v1/chat/completions",
            json=payload,
            timeout=self.timeout
        )
        
        if response.status_code == 200:
            result = response.json()
            return result.get('choices', [{}])[0].get('message', {}).get('content', '').strip()
        else:
            raise Exception(f"LM Studio API error: {response.status_code} - {response.text}")
    
    def is_available(self) -> bool:
        """Check if AI provider is available"""
        try:
            if self.provider == "openai":
                response = self.session.get(f"{self.base_url}/models", timeout=5)
            elif self.provider == "anthropic":
                response = self.session.get(f"{self.base_url}/messages", timeout=5)
            elif self.provider == "grok":
                response = self.session.get(f"{self.base_url}/models", timeout=5)
            elif self.provider == "google":
                response = self.session.get(f"{self.base_url}/models?key={self.api_key}", timeout=5)
            else:  # LM Studio
                response = self.session.get(f"{self.base_url}/v1/models", timeout=5)
            
            return response.status_code == 200
        except:
            return False
    
    def wait_for_model_ready(self, timeout: int = 60) -> bool:
        """Wait for the model to be ready to accept requests"""
        print("Waiting for model to load...")
        start_time = time.time()
        
        while time.time() - start_time < timeout:
            try:
                # Try a simple test request
                test_response = self.generate_text("Hello", max_tokens=5, temperature=0.1)
                if test_response and test_response.strip():
                    print("✓ Model is ready")
                    return True
            except:
                pass
            
            print(".", end="", flush=True)
            time.sleep(2)
        
        print(f"\n⚠ Model not ready after {timeout} seconds")
        return False
    
    def get_available_models(self) -> List[str]:
        """Get list of available models from LM Studio"""
        try:
            response = self.session.get(f"{self.base_url}/v1/models", timeout=10)
            if response.status_code == 200:
                models_data = response.json()
                models = []
                for model in models_data.get('data', []):
                    model_id = model.get('id', '')
                    if model_id:
                        models.append(model_id)
                return models
            return []
        except Exception as e:
            print(f"Error fetching models: {e}")
            return []
    
    def is_model_available(self, model_name: str) -> bool:
        """Check if a specific model is available"""
        available_models = self.get_available_models()
        return model_name in available_models
    
    def get_model_info(self) -> Dict:
        """Get detailed information about the current model"""
        try:
            response = self.session.get(f"{self.base_url}/v1/models", timeout=10)
            if response.status_code == 200:
                models_data = response.json()
                for model in models_data.get('data', []):
                    if model.get('id') == self.model:
                        return model
            return {}
        except Exception as e:
            print(f"Error fetching model info: {e}")
            return {}
    
    def get_context_window(self) -> int:
        """Get the context window size for the current model"""
        model_info = self.get_model_info()
        # Try different possible fields for context window
        context_fields = ['context_length', 'max_tokens', 'max_context_length', 'context_window']
        for field in context_fields:
            if field in model_info:
                return int(model_info[field])
        
        # Default fallbacks based on common model sizes
        if 'gpt-oss-20b' in self.model:
            return 8192
        elif 'gpt-oss-7b' in self.model:
            return 4096
        elif 'llama-2' in self.model:
            return 4096
        elif 'mistral' in self.model:
            return 32768
        else:
            return 4096  # Conservative default
    
    def estimate_tokens(self, text: str) -> int:
        """Rough estimation of token count (4 chars per token average)"""
        return len(text) // 4


def select_lm_studio_model(lm_client: AIClient, preferred_model: str = "qwen/qwen3-4b-2507") -> str:
    """Select an available model from LM Studio, with fallback options"""
    if not lm_client.is_available():
        return None
    
    available_models = lm_client.get_available_models()
    if not available_models:
        print("No models found in LM Studio")
        return None
    
    # Check if preferred model is available
    if preferred_model in available_models:
        print(f"✓ Using preferred model: {preferred_model}")
        # Wait for model to be ready
        if lm_client.wait_for_model_ready():
            return preferred_model
        else:
            print("⚠ Model failed to load, trying fallback...")
    else:
        print(f"⚠ Preferred model '{preferred_model}' not available")
        print("Available models:")
        for i, model in enumerate(available_models, 1):
            print(f"  {i}. {model}")
    
    # Try to find a reasonable fallback
    fallback_models = [
        "qwen/qwen3-4b-2507", "qwen2.5-0.5b-instruct-mlx", "qwen3-coder-30b-a3b-instruct-mlx",
        "gpt-oss-20b", "gpt-oss-7b", "gpt-oss-3b", "gpt-oss-1b",
        "llama-2-70b", "llama-2-13b", "llama-2-7b",
        "codellama-34b", "codellama-13b", "codellama-7b",
        "mistral-7b", "mixtral-8x7b"
    ]
    
    for fallback in fallback_models:
        if fallback in available_models:
            print(f"✓ Trying fallback model: {fallback}")
            lm_client.model = fallback
            if lm_client.wait_for_model_ready():
                return fallback
            else:
                print(f"⚠ {fallback} failed to load, trying next...")
    
    # If no fallback found, use the first available model
    selected_model = available_models[0]
    print(f"✓ Trying first available model: {selected_model}")
    lm_client.model = selected_model
    if lm_client.wait_for_model_ready():
        return selected_model
    
    print("⚠ No models could be loaded successfully")
    return None


@dataclass
class SubtitleBlock:
    """Represents a single subtitle block from SRT file"""
    index: int
    start_time: str
    end_time: str
    text: str
    start_seconds: float
    end_seconds: float


@dataclass
class Paragraph:
    """Represents a grouped paragraph with metadata"""
    start_time: str
    start_seconds: float
    text: str
    subtitle_blocks: List[SubtitleBlock]


class SubtitleParser:
    """Parses SRT and VTT subtitle files"""
    
    def __init__(self):
        # SRT time format: 00:00:00,000
        self.srt_time_pattern = re.compile(r'(\d{2}):(\d{2}):(\d{2}),(\d{3})')
        # VTT time format: 00:00:00.000
        self.vtt_time_pattern = re.compile(r'(\d{2}):(\d{2}):(\d{2})\.(\d{3})')
    
    def parse_time_to_seconds(self, time_str: str) -> float:
        """Convert various time formats to seconds"""
        # Try SRT format first (comma separator)
        match = self.srt_time_pattern.match(time_str)
        if match:
            hours, minutes, seconds, milliseconds = map(int, match.groups())
            return hours * 3600 + minutes * 60 + seconds + milliseconds / 1000.0
        
        # Try VTT format (dot separator)
        match = self.vtt_time_pattern.match(time_str)
        if match:
            hours, minutes, seconds, milliseconds = map(int, match.groups())
            return hours * 3600 + minutes * 60 + seconds + milliseconds / 1000.0
        
        # Try ASS/SSA format (H:MM:SS.CC - centiseconds)
        ass_pattern = re.compile(r'(\d+):(\d{2}):(\d{2})\.(\d{2})')
        match = ass_pattern.match(time_str)
        if match:
            hours, minutes, seconds, centiseconds = map(int, match.groups())
            return hours * 3600 + minutes * 60 + seconds + centiseconds / 100.0
        
        # Try SBV format (H:MM:SS.mmm)
        sbv_pattern = re.compile(r'(\d+):(\d{2}):(\d{2})\.(\d{3})')
        match = sbv_pattern.match(time_str)
        if match:
            hours, minutes, seconds, milliseconds = map(int, match.groups())
            return hours * 3600 + minutes * 60 + seconds + milliseconds / 1000.0
        
        # Try LRC format ([MM:SS.mm] or [H:MM:SS.mm])
        lrc_pattern = re.compile(r'\[(\d+):(\d{2})(?:\.(\d{2}))?\]')
        match = lrc_pattern.match(time_str)
        if match:
            minutes, seconds, centiseconds = match.groups()
            minutes = int(minutes)
            seconds = int(seconds)
            centiseconds = int(centiseconds) if centiseconds else 0
            return minutes * 60 + seconds + centiseconds / 100.0
        
        return 0.0
    
    def detect_format(self, file_path: str) -> str:
        """Detect subtitle format based on file extension and content"""
        # Check file extension first
        file_ext = file_path.lower().split('.')[-1] if '.' in file_path else ''
        
        # Extension-based detection
        if file_ext == 'vtt':
            return 'vtt'
        elif file_ext == 'srt':
            return 'srt'
        elif file_ext == 'ass':
            return 'ass'
        elif file_ext == 'ssa':
            return 'ssa'
        elif file_ext == 'sbv':
            return 'sbv'
        elif file_ext == 'ttml' or file_ext == 'dfxp':
            return 'ttml'
        elif file_ext == 'smi':
            return 'sami'
        elif file_ext == 'lrc':
            return 'lrc'
        
        # If no extension or unknown, check content
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read(1000)  # Read first 1000 chars for analysis
                content_lower = content.lower()
                
                # Check for format signatures
                if 'webvtt' in content_lower:
                    return 'vtt'
                elif '[script info]' in content_lower or '[v4+ styles]' in content_lower:
                    return 'ass'
                elif '[script info]' in content_lower and 'ssa' in content_lower:
                    return 'ssa'
                elif '<tt' in content_lower and 'xml' in content_lower:
                    return 'ttml'
                elif '<sami>' in content_lower or '<sync' in content_lower:
                    return 'sami'
                elif '[' in content and ']' in content and ':' in content:
                    # Check if it looks like LRC format
                    lines = content.split('\n')[:3]
                    if all('[' in line and ']' in line for line in lines if line.strip()):
                        return 'lrc'
                elif content.split('\n')[0].strip().isdigit():
                    return 'srt'
        except:
            pass
        
        # Default to SRT for backward compatibility
        return 'srt'
    
    def parse_srt_file(self, file_path: str) -> List[SubtitleBlock]:
        """Parse SRT or VTT file and return list of subtitle blocks with enhanced error handling"""
        blocks = []
        
        try:
            # Detect file format
            file_format = self.detect_format(file_path)
            log(f"Detected format: {file_format.upper()}")
            
            # Try multiple encodings
            encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252']
            content = None
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        content = f.read()
                    log(f"Successfully read file with {encoding} encoding")
                    break
                except UnicodeDecodeError:
                    continue
            
            if content is None:
                raise ValueError("Could not read file with any supported encoding")
            
            # Parse based on format
            if file_format == 'vtt':
                subtitle_blocks = self._parse_vtt_content(content)
            elif file_format == 'ass':
                subtitle_blocks = self._parse_ass_content(content)
            elif file_format == 'ssa':
                subtitle_blocks = self._parse_ssa_content(content)
            elif file_format == 'sbv':
                subtitle_blocks = self._parse_sbv_content(content)
            elif file_format == 'ttml':
                subtitle_blocks = self._parse_ttml_content(content)
            elif file_format == 'sami':
                subtitle_blocks = self._parse_sami_content(content)
            elif file_format == 'lrc':
                subtitle_blocks = self._parse_lrc_content(content)
            else:
                subtitle_blocks = self._parse_srt_content(content)
            
            log(f"Found {len(subtitle_blocks)} subtitle blocks")
            
            for i, block in enumerate(subtitle_blocks):
                try:
                    lines = block.strip().split('\n')
                    if len(lines) < 3:
                        log(f"Skipping malformed block {i+1}: insufficient lines")
                        continue
                    
                    # Handle VTT format differences
                    if file_format == 'vtt':
                        # VTT blocks: timestamp line, text lines (no index)
                        time_line = lines[0]
                        text = '\n'.join(lines[1:])
                        index = i + 1  # Generate index for VTT
                    else:
                        # SRT blocks: index, timestamp line, text lines
                        index = int(lines[0])
                        time_line = lines[1]
                        text = '\n'.join(lines[2:])
                    
                    # Parse time range
                    if ' --> ' not in time_line:
                        log(f"Skipping block {i+1}: invalid timestamp format")
                        continue
                    
                    start_time, end_time = time_line.split(' --> ')
                    start_seconds = self.parse_time_to_seconds(start_time)
                    end_seconds = self.parse_time_to_seconds(end_time)
                    
                    blocks.append(SubtitleBlock(
                        index=index,
                        start_time=start_time,
                        end_time=end_time,
                        text=text.strip(),
                        start_seconds=start_seconds,
                        end_seconds=end_seconds
                    ))
                    
                except (ValueError, IndexError) as e:
                    log(f"Error parsing block {i+1}: {e}")
                    continue
                except Exception as e:
                    log(f"Unexpected error in block {i+1}: {e}")
                    continue
            
            log(f"Successfully parsed {len(blocks)} subtitle blocks")
            return blocks
            
        except FileNotFoundError:
            log(f"File not found: {file_path}", "ERROR")
            raise
        except Exception as e:
            log(f"Error parsing SRT file: {e}", "ERROR")
            raise
    
    def _parse_srt_content(self, content: str) -> List[str]:
        """Parse SRT content into subtitle blocks"""
        return content.strip().split('\n\n')
    
    def _parse_vtt_content(self, content: str) -> List[str]:
        """Parse VTT content into subtitle blocks"""
        lines = content.strip().split('\n')
        blocks = []
        current_block = []
        
        for line in lines:
            line = line.strip()
            
            # Skip WEBVTT header and empty lines
            if line == 'WEBVTT' or line == '':
                continue
            
            # Skip style and note blocks
            if line.startswith('NOTE') or line.startswith('STYLE'):
                continue
            
            # Check if this is a timestamp line (contains -->)
            if ' --> ' in line:
                # If we have a current block, save it
                if current_block:
                    blocks.append('\n'.join(current_block))
                    current_block = []
                # Start new block with timestamp
                current_block = [line]
            else:
                # This is text content
                if current_block:  # Only add if we're in a block
                    current_block.append(line)
        
        # Add the last block if it exists
        if current_block:
            blocks.append('\n'.join(current_block))
        
        return blocks
    
    def _parse_ass_content(self, content: str) -> List[str]:
        """Parse ASS content into subtitle blocks"""
        blocks = []
        lines = content.split('\n')
        current_block = []
        in_dialogue = False
        
        for line in lines:
            line = line.strip()
            if line.startswith('[Events]'):
                in_dialogue = True
                continue
            elif line.startswith('[') and line.endswith(']'):
                in_dialogue = False
                continue
            
            if in_dialogue and line.startswith('Dialogue:'):
                # ASS format: Dialogue: Layer,Start,End,Style,Name,MarginL,MarginR,MarginV,Effect,Text
                parts = line.split(',', 9)
                if len(parts) >= 10:
                    start_time = parts[1]
                    end_time = parts[2]
                    text = parts[9]
                    # Clean up ASS formatting codes
                    text = re.sub(r'\{[^}]*\}', '', text)  # Remove style codes
                    text = text.replace('\\N', '\n')  # Convert line breaks
                    text = text.replace('\\n', '\n')
                    blocks.append(f"{start_time} --> {end_time}\n{text}")
        
        return blocks
    
    def _parse_ssa_content(self, content: str) -> List[str]:
        """Parse SSA content into subtitle blocks (similar to ASS)"""
        return self._parse_ass_content(content)
    
    def _parse_sbv_content(self, content: str) -> List[str]:
        """Parse SBV content into subtitle blocks"""
        blocks = []
        lines = content.split('\n')
        current_block = []
        
        for line in lines:
            line = line.strip()
            if not line:
                if current_block:
                    blocks.append('\n'.join(current_block))
                    current_block = []
            elif ',' in line and ':' in line:
                # SBV format: start_time,end_time
                current_block = [line]
            else:
                if current_block:
                    current_block.append(line)
        
        if current_block:
            blocks.append('\n'.join(current_block))
        
        return blocks
    
    def _parse_ttml_content(self, content: str) -> List[str]:
        """Parse TTML content into subtitle blocks"""
        blocks = []
        try:
            import xml.etree.ElementTree as ET
            root = ET.fromstring(content)
            
            # Find all p elements (paragraphs)
            for p in root.findall('.//{http://www.w3.org/ns/ttml}p'):
                begin = p.get('begin')
                end = p.get('end')
                text = p.text or ''
                
                # Clean up text
                text = re.sub(r'<[^>]+>', '', text)  # Remove XML tags
                text = text.strip()
                
                if begin and end and text:
                    # Convert TTML time format to SRT-like format
                    start_time = self._convert_ttml_time(begin)
                    end_time = self._convert_ttml_time(end)
                    blocks.append(f"{start_time} --> {end_time}\n{text}")
        except Exception as e:
            log(f"Error parsing TTML: {e}")
        
        return blocks
    
    def _convert_ttml_time(self, ttml_time: str) -> str:
        """Convert TTML time format to SRT format"""
        # TTML can have formats like: 1.5s, 00:00:01.500, 1.5
        if ttml_time.endswith('s'):
            seconds = float(ttml_time[:-1])
        elif ':' in ttml_time:
            # Already in HH:MM:SS.mmm format
            return ttml_time.replace('.', ',')
        else:
            seconds = float(ttml_time)
        
        # Convert to HH:MM:SS,mmm format
        hours = int(seconds // 3600)
        minutes = int((seconds % 3600) // 60)
        secs = int(seconds % 60)
        millis = int((seconds % 1) * 1000)
        
        return f"{hours:02d}:{minutes:02d}:{secs:02d},{millis:03d}"
    
    def _parse_sami_content(self, content: str) -> List[str]:
        """Parse SAMI content into subtitle blocks"""
        blocks = []
        try:
            import xml.etree.ElementTree as ET
            root = ET.fromstring(content)
            
            for sync in root.findall('.//sync'):
                start = sync.get('start')
                text_elem = sync.find('p')
                if text_elem is not None:
                    text = text_elem.text or ''
                    text = re.sub(r'<[^>]+>', '', text)  # Remove HTML tags
                    text = text.strip()
                    
                    if start and text:
                        # Convert milliseconds to time format
                        start_ms = int(start)
                        end_ms = start_ms + 3000  # Default 3 seconds duration
                        
                        start_time = self._convert_ms_to_time(start_ms)
                        end_time = self._convert_ms_to_time(end_ms)
                        
                        blocks.append(f"{start_time} --> {end_time}\n{text}")
        except Exception as e:
            log(f"Error parsing SAMI: {e}")
        
        return blocks
    
    def _convert_ms_to_time(self, milliseconds: int) -> str:
        """Convert milliseconds to HH:MM:SS,mmm format"""
        seconds = milliseconds // 1000
        millis = milliseconds % 1000
        
        hours = seconds // 3600
        minutes = (seconds % 3600) // 60
        secs = seconds % 60
        
        return f"{hours:02d}:{minutes:02d}:{secs:02d},{millis:03d}"
    
    def _parse_lrc_content(self, content: str) -> List[str]:
        """Parse LRC content into subtitle blocks"""
        blocks = []
        lines = content.split('\n')
        
        for i, line in enumerate(lines):
            line = line.strip()
            if not line or not line.startswith('['):
                continue
            
            # Extract timestamp and text
            match = re.match(r'\[(\d+:\d{2}(?:\.\d{2})?)\](.*)', line)
            if match:
                time_str = match.group(1)
                text = match.group(2).strip()
                
                if text:
                    # Convert LRC time to seconds, then to SRT format
                    seconds = self.parse_time_to_seconds(f"[{time_str}]")
                    end_seconds = seconds + 3.0  # Default 3 seconds duration
                    
                    start_time = self._convert_seconds_to_srt_time(seconds)
                    end_time = self._convert_seconds_to_srt_time(end_seconds)
                    
                    blocks.append(f"{start_time} --> {end_time}\n{text}")
        
        return blocks
    
    def _convert_seconds_to_srt_time(self, seconds: float) -> str:
        """Convert seconds to SRT time format"""
        hours = int(seconds // 3600)
        minutes = int((seconds % 3600) // 60)
        secs = int(seconds % 60)
        millis = int((seconds % 1) * 1000)
        
        return f"{hours:02d}:{minutes:02d}:{secs:02d},{millis:03d}"


class ParagraphGrouper:
    """Groups subtitle blocks into logical paragraphs using AI assistance"""
    
    def __init__(self, max_gap_seconds: float = 3.0, min_paragraph_length: int = 10, lm_client: Optional[AIClient] = None):
        self.max_gap_seconds = max_gap_seconds
        self.min_paragraph_length = min_paragraph_length
        self.lm_client = lm_client
    
    def group_into_paragraphs(self, blocks: List[SubtitleBlock]) -> List[Paragraph]:
        """Group subtitle blocks into paragraphs based on timing and content"""
        if not blocks:
            return []
        
        # Use AI for paragraph detection if available
        if self.lm_client and self.lm_client.is_available():
            log("Using AI-enhanced paragraph grouping...")
            return self._group_with_ai_chunking(blocks)
        
        # For very long files, use smart content analysis
        if len(blocks) > 1000:
            print("Using smart content analysis for long file...")
            return self._group_with_smart_chunking(blocks)
        
        return self._group_regular_paragraphs(blocks)
    
    def _group_regular_paragraphs(self, blocks: List[SubtitleBlock]) -> List[Paragraph]:
        """Regular paragraph grouping based on timing and content with sentence boundary respect"""
        import re
        
        paragraphs = []
        current_paragraph_blocks = [blocks[0]]
        
        for i in range(1, len(blocks)):
            current_block = blocks[i]
            previous_block = blocks[i-1]
            
            # Calculate gap between subtitles
            gap = current_block.start_seconds - previous_block.end_seconds
            
            # Check if we should start a new paragraph
            should_break = (
                gap > self.max_gap_seconds or
                self._is_speaker_change(previous_block.text, current_block.text) or
                self._is_topic_change(previous_block.text, current_block.text)
            )
            
            if should_break:
                # Check if this is a good sentence boundary
                if self._is_good_sentence_boundary(blocks, i):
                    # Finalize current paragraph
                    if current_paragraph_blocks:
                        paragraph = self._create_paragraph(current_paragraph_blocks)
                        if paragraph and len(paragraph.text) >= self.min_paragraph_length:
                            paragraphs.append(paragraph)
                        current_paragraph_blocks = [current_block]
                    else:
                        current_paragraph_blocks = [current_block]
                else:
                    # Not a good boundary, continue building current paragraph
                    current_paragraph_blocks.append(current_block)
            else:
                current_paragraph_blocks.append(current_block)
        
        # Add final paragraph
        if current_paragraph_blocks:
            paragraph = self._create_paragraph(current_paragraph_blocks)
            if paragraph and len(paragraph.text) >= self.min_paragraph_length:
                paragraphs.append(paragraph)
        
        return paragraphs
    
    def _group_with_ai_chunking(self, blocks: List[SubtitleBlock]) -> List[Paragraph]:
        """Use AI-assisted chunking for very long subtitle files"""
        log("Using AI-assisted paragraph grouping for long file...")
        
        # Process in smaller chunks for better analysis
        chunk_size = 100
        all_paragraphs = []
        processed_blocks = set()
        
        # Calculate number of chunks for progress bar
        num_chunks = len(range(0, len(blocks), chunk_size // 2))
        
        with tqdm(total=num_chunks, desc="Processing chunks", unit="chunk") as pbar:
            for start_idx in range(0, len(blocks), chunk_size // 2):  # 50% overlap
                end_idx = min(start_idx + chunk_size, len(blocks))
                chunk = blocks[start_idx:end_idx]
                
                if not chunk:
                    pbar.update(1)
                    continue
                
                log(f"Processing blocks {start_idx+1}-{end_idx} with AI...")
                
                # Use AI analysis for this chunk
                chunk_paragraphs = self._ai_find_natural_boundaries(chunk, start_idx)
                
                # Only add paragraphs that haven't been processed yet
                for paragraph in chunk_paragraphs:
                    paragraph_start_idx = start_idx + blocks.index(paragraph.subtitle_blocks[0])
                    if paragraph_start_idx not in processed_blocks:
                        all_paragraphs.append(paragraph)
                        for block in paragraph.subtitle_blocks:
                            block_idx = blocks.index(block)
                            processed_blocks.add(block_idx)
                
                pbar.update(1)
        
        # Sort paragraphs by their start time
        all_paragraphs.sort(key=lambda p: p.start_seconds)
        
        log(f"Created {len(all_paragraphs)} paragraphs using AI")
        return all_paragraphs
    
    def _group_with_smart_chunking(self, blocks: List[SubtitleBlock]) -> List[Paragraph]:
        """Use fallback rule-based chunking for very long subtitle files"""
        print("Using fallback paragraph grouping for long file...")
        
        # Process in smaller chunks for better analysis
        chunk_size = 100
        all_paragraphs = []
        processed_blocks = set()
        
        for start_idx in range(0, len(blocks), chunk_size // 2):  # 50% overlap
            end_idx = min(start_idx + chunk_size, len(blocks))
            chunk = blocks[start_idx:end_idx]
            
            if not chunk:
                continue
            
            print(f"Processing blocks {start_idx+1}-{end_idx} with fallback rules...")
            
            # Use fallback analysis for this chunk
            chunk_paragraphs = self._ai_find_natural_boundaries(chunk, start_idx)
            
            # Only add paragraphs that haven't been processed yet
            for paragraph in chunk_paragraphs:
                paragraph_start_idx = start_idx + blocks.index(paragraph.subtitle_blocks[0])
                if paragraph_start_idx not in processed_blocks:
                    all_paragraphs.append(paragraph)
                    for block in paragraph.subtitle_blocks:
                        block_idx = blocks.index(block)
                        processed_blocks.add(block_idx)
        
        # Sort paragraphs by their start time
        all_paragraphs.sort(key=lambda p: p.start_seconds)
        
        print(f"Created {len(all_paragraphs)} paragraphs using fallback rules")
        return all_paragraphs
    
    def _ai_find_natural_boundaries(self, blocks: List[SubtitleBlock], start_offset: int = 0) -> List[Paragraph]:
        """Use AI to find natural paragraph boundaries"""
        print(f"  Using AI analysis for {len(blocks)} blocks")
        
        if not self.lm_client or not self.lm_client.is_available():
            print("  AI not available, using fallback rules")
            return self._fallback_find_boundaries(blocks)
        
        # Prepare text for AI analysis
        text_lines = []
        for i, block in enumerate(blocks):
            text_lines.append(f"{i}: {block.text}")
        
        # Create an improved prompt for the AI that respects sentence boundaries
        prompt = f"""Analyze this transcript and find natural paragraph breaks. Consider:
1. Complete sentences - don't break mid-sentence
2. Natural conversation flow
3. Topic changes
4. Speaker changes
5. Pauses in speech

Transcript:
{chr(10).join(text_lines)}

Return ONLY the line numbers where paragraphs should end (after complete sentences), separated by commas:"""
        
        try:
            response = self.lm_client.generate_text(prompt, max_tokens=100, temperature=0.1)
            if response and response.strip():
                print(f"  AI response: {response.strip()}")
                # Parse the response to get break points
                break_indices = []
                for part in response.strip().split(','):
                    try:
                        break_point = int(part.strip())
                        if 0 <= break_point < len(blocks):
                            break_indices.append(break_point)
                    except ValueError:
                        continue
                
                if break_indices:
                    print(f"  AI found {len(break_indices)} break points: {break_indices[:5]}")
                    # Post-process to ensure we don't break mid-sentence
                    validated_breaks = self._validate_sentence_boundaries(blocks, break_indices)
                    return self._create_paragraphs_from_breaks(blocks, validated_breaks)
        except Exception as e:
            print(f"  AI paragraph detection failed: {e}")
        
        # Fallback to rule-based approach
        print("  Falling back to rule-based analysis")
        return self._fallback_find_boundaries(blocks)
    
    def _fallback_find_boundaries(self, blocks: List[SubtitleBlock]) -> List[Paragraph]:
        """Fallback rule-based paragraph detection with sentence boundary respect"""
        import re
        
        break_indices = []
        
        for i in range(1, len(blocks)):
            current_block = blocks[i]
            previous_block = blocks[i-1]
            
            should_break = False
            
            # 1. Questions (new paragraph after questions)
            if previous_block.text.strip().endswith('?'):
                should_break = True
            
            # 2. Time gaps (more than 3 seconds)
            elif current_block.start_seconds - previous_block.end_seconds > 3.0:
                should_break = True
            
            # 3. Short responses followed by longer content
            elif (len(previous_block.text.strip()) < 15 and 
                  len(current_block.text.strip()) > 40):
                should_break = True
            
            # 4. Conversation starters
            elif any(starter in current_block.text.lower() for starter in [
                'so,', 'well,', 'now,', 'okay,', 'right,', 'and so', 'and then',
                'let me', 'i think', 'i believe', 'i would', 'i can'
            ]):
                should_break = True
            
            # 5. Topic change indicators
            elif any(indicator in current_block.text.lower() for indicator in [
                'speaking of', 'on the topic of', 'regarding', 'about',
                'in terms of', 'when it comes to', 'as for'
            ]):
                should_break = True
            
            # 6. Every 8-12 blocks as fallback (aim for 5-10 paragraphs per 100 blocks)
            elif i % 10 == 0:
                should_break = True
            
            if should_break:
                # Validate that this is a good sentence boundary
                if self._is_good_sentence_boundary(blocks, i):
                    break_indices.append(i)
                else:
                    # Try to find a better boundary nearby
                    better_break = self._find_nearby_sentence_boundary(blocks, i)
                    if better_break is not None and better_break not in break_indices:
                        break_indices.append(better_break)
        
        # Limit to reasonable number of breaks
        if len(break_indices) > len(blocks) // 8:
            # Too many breaks, keep only every few
            break_indices = break_indices[::max(1, len(break_indices) // 8)]
        elif len(break_indices) < 3:
            # Too few breaks, add some
            break_indices = list(range(10, len(blocks), 12))
        
        print(f"  Found {len(break_indices)} natural break points: {break_indices[:5]}")
        return self._create_paragraphs_from_breaks(blocks, break_indices)
    
    def _is_good_sentence_boundary(self, blocks: List[SubtitleBlock], break_idx: int) -> bool:
        """Check if a break point is at a good sentence boundary"""
        import re
        
        if break_idx >= len(blocks):
            return False
            
        current_block = blocks[break_idx]
        next_block = blocks[break_idx + 1] if break_idx + 1 < len(blocks) else None
        
        current_text = current_block.text.strip()
        next_text = next_block.text.strip() if next_block else ""
        
        # Good boundary if current block ends with sentence punctuation
        if re.search(r'[.!?]$', current_text):
            return True
        
        # Good boundary if next block starts with capital letter
        if next_text and re.match(r'^[A-Z]', next_text):
            return True
        
        return False
    
    def _find_nearby_sentence_boundary(self, blocks: List[SubtitleBlock], break_idx: int) -> Optional[int]:
        """Find a nearby sentence boundary within 2 blocks"""
        import re
        
        # Look ahead up to 2 blocks
        for i in range(break_idx + 1, min(break_idx + 3, len(blocks))):
            block_text = blocks[i].text.strip()
            
            # Check if this block ends with sentence punctuation
            if re.search(r'[.!?]$', block_text):
                return i
            
            # Check if next block starts with capital letter
            if i + 1 < len(blocks):
                next_text = blocks[i + 1].text.strip()
                if re.match(r'^[A-Z]', next_text):
                    return i
        
        # If no good boundary found, use the original
        return break_idx
    
    def _create_paragraphs_from_breaks(self, blocks: List[SubtitleBlock], break_indices: List[int]) -> List[Paragraph]:
        """Create paragraphs from break indices"""
        paragraphs = []
        current_blocks = []
        
        for i, block in enumerate(blocks):
            current_blocks.append(block)
            
            # Check if this is a break point
            if i in break_indices or i == len(blocks) - 1:
                if current_blocks:
                    paragraph = self._create_paragraph(current_blocks)
                    if paragraph and len(paragraph.text) >= self.min_paragraph_length:
                        paragraphs.append(paragraph)
                    current_blocks = []
        
        print(f"  Created {len(paragraphs)} paragraphs from {len(blocks)} blocks")
        return paragraphs
    
    def _validate_sentence_boundaries(self, blocks: List[SubtitleBlock], break_indices: List[int]) -> List[int]:
        """Validate and adjust break points to respect sentence boundaries"""
        import re
        
        validated_breaks = []
        
        for break_idx in break_indices:
            if break_idx >= len(blocks):
                continue
                
            # Check if the break point is at a natural sentence boundary
            current_block = blocks[break_idx]
            next_block = blocks[break_idx + 1] if break_idx + 1 < len(blocks) else None
            
            # Look for sentence endings in current block
            current_text = current_block.text.strip()
            next_text = next_block.text.strip() if next_block else ""
            
            # Check if current block ends with sentence punctuation
            if re.search(r'[.!?]$', current_text):
                validated_breaks.append(break_idx)
                continue
            
            # Check if next block starts with capital letter (new sentence)
            if next_text and re.match(r'^[A-Z]', next_text):
                validated_breaks.append(break_idx)
                continue
            
            # If not a natural boundary, try to find the next sentence boundary
            adjusted_break = self._find_next_sentence_boundary(blocks, break_idx)
            if adjusted_break is not None and adjusted_break not in validated_breaks:
                validated_breaks.append(adjusted_break)
        
        return sorted(set(validated_breaks))
    
    def _find_next_sentence_boundary(self, blocks: List[SubtitleBlock], start_idx: int) -> Optional[int]:
        """Find the next natural sentence boundary after start_idx"""
        import re
        
        # Look ahead up to 3 blocks for a sentence boundary
        for i in range(start_idx + 1, min(start_idx + 4, len(blocks))):
            block_text = blocks[i].text.strip()
            
            # Check if this block ends with sentence punctuation
            if re.search(r'[.!?]$', block_text):
                return i
            
            # Check if next block starts with capital letter
            if i + 1 < len(blocks):
                next_text = blocks[i + 1].text.strip()
                if re.match(r'^[A-Z]', next_text):
                    return i
        
        # If no natural boundary found, use the original break point
        return start_idx
    
    def _group_chunk_regular(self, blocks: List[SubtitleBlock]) -> List[Paragraph]:
        """Regular grouping for a chunk of blocks"""
        paragraphs = []
        current_paragraph_blocks = [blocks[0]] if blocks else []
        
        for i in range(1, len(blocks)):
            current_block = blocks[i]
            previous_block = blocks[i-1]
            
            gap = current_block.start_seconds - previous_block.end_seconds
            
            should_break = (
                gap > self.max_gap_seconds or
                self._is_speaker_change(previous_block.text, current_block.text) or
                self._is_topic_change(previous_block.text, current_block.text)
            )
            
            if should_break:
                if current_paragraph_blocks:
                    paragraph = self._create_paragraph(current_paragraph_blocks)
                    if paragraph and len(paragraph.text) >= self.min_paragraph_length:
                        paragraphs.append(paragraph)
                    current_paragraph_blocks = [current_block]
                else:
                    current_paragraph_blocks = [current_block]
            else:
                current_paragraph_blocks.append(current_block)
        
        if current_paragraph_blocks:
            paragraph = self._create_paragraph(current_paragraph_blocks)
            if paragraph and len(paragraph.text) >= self.min_paragraph_length:
                paragraphs.append(paragraph)
        
        return paragraphs
    
    def _create_paragraph(self, blocks: List[SubtitleBlock]) -> Optional[Paragraph]:
        """Create a paragraph from subtitle blocks"""
        if not blocks:
            return None
        
        # Combine text from all blocks
        text_parts = []
        for block in blocks:
            # Clean up text (remove HTML tags, normalize whitespace)
            clean_text = re.sub(r'<[^>]+>', '', block.text)
            clean_text = re.sub(r'\s+', ' ', clean_text).strip()
            if clean_text:
                text_parts.append(clean_text)
        
        if not text_parts:
            return None
        
        combined_text = ' '.join(text_parts)
        
        return Paragraph(
            start_time=blocks[0].start_time,
            start_seconds=blocks[0].start_seconds,
            text=combined_text,
            subtitle_blocks=blocks
        )
    
    def _is_speaker_change(self, text1: str, text2: str) -> bool:
        """Detect potential speaker changes"""
        # Look for patterns that might indicate speaker changes
        speaker_indicators = [
            r'^[A-Z][a-z]+:',  # "John:"
            r'^[A-Z][A-Z]+:',  # "JOHN:"
            r'^[A-Z][a-z]+\s+[A-Z][a-z]+:',  # "John Smith:"
        ]
        
        for pattern in speaker_indicators:
            if re.search(pattern, text2):
                return True
        
        return False
    
    def _is_topic_change(self, text1: str, text2: str) -> bool:
        """Detect potential topic changes"""
        # Simple heuristics for topic changes
        topic_indicators = [
            r'^(So|Now|Next|Then|Also|However|But|And|Well)',
            r'^(Let\'s|Let us)',
            r'^(Moving on|Speaking of)',
        ]
        
        for pattern in topic_indicators:
            if re.search(pattern, text2, re.IGNORECASE):
                return True
        
        return False


class HeadlineGenerator:
    """Generates headlines and subheadlines using AI analysis"""
    
    def __init__(self, lm_client: Optional[AIClient] = None):
        self.lm_client = lm_client
        self.headline_keywords = [
            'introduction', 'overview', 'summary', 'conclusion',
            'background', 'context', 'setup', 'preparation',
            'main', 'primary', 'key', 'important', 'critical',
            'discussion', 'analysis', 'review', 'examination',
            'next', 'following', 'subsequent', 'additional'
        ]
    
    def generate_headlines(self, paragraphs: List[Paragraph]) -> List[Tuple[int, str, str]]:
        """Generate headlines for paragraphs using AI analysis"""
        headlines = []
        
        # Use AI for headline generation if available
        if self.lm_client and self.lm_client.is_available():
            headlines = self._generate_ai_headlines(paragraphs)
        else:
            # Fallback to rule-based generation - only for major sections
            headlines = []
            
            # Always include first paragraph as introduction
            if paragraphs:
                headlines.append((0, "Introduction", self._get_timestamp_marker(paragraphs[0].start_time)))
            
            # Look for major topic changes
            for i in range(1, len(paragraphs)):
                current_para = paragraphs[i]
                previous_para = paragraphs[i-1]
                
                topic_change = self._rule_based_topic_change(previous_para, current_para)
                if topic_change:
                    headlines.append((i, topic_change, self._get_timestamp_marker(current_para.start_time)))
            
            # Always include last paragraph as conclusion if there are multiple paragraphs
            if len(paragraphs) > 1:
                last_idx = len(paragraphs) - 1
                headlines.append((last_idx, "Conclusion", self._get_timestamp_marker(paragraphs[last_idx].start_time)))
        
        return headlines
    
    def _generate_ai_headlines(self, paragraphs: List[Paragraph]) -> List[Tuple[int, str, str]]:
        """Generate headlines using AI analysis - only for major sections"""
        headlines = []
        
        # Only generate headlines for significant sections, not every paragraph
        # Look for topic changes, introductions, conclusions, and major transitions
        
        # Always include first paragraph as introduction
        if paragraphs:
            headlines.append((0, "Introduction", self._get_timestamp_marker(paragraphs[0].start_time)))
        
        # Analyze for major topic changes
        for i in range(1, len(paragraphs)):
            current_para = paragraphs[i]
            previous_para = paragraphs[i-1]
            
            # Check for major topic changes using AI
            topic_change = self._detect_topic_change(previous_para, current_para)
            if topic_change:
                headlines.append((i, topic_change, self._get_timestamp_marker(current_para.start_time)))
        
        # Always include last paragraph as conclusion if there are multiple paragraphs
        if len(paragraphs) > 1:
            last_idx = len(paragraphs) - 1
            headlines.append((last_idx, "Conclusion", self._get_timestamp_marker(paragraphs[last_idx].start_time)))
        
        return headlines
    
    def _detect_topic_change(self, prev_para: Paragraph, current_para: Paragraph) -> Optional[str]:
        """Detect if there's a significant topic change between paragraphs"""
        if not self.lm_client or not self.lm_client.is_available():
            return self._rule_based_topic_change(prev_para, current_para)
        
        # Use AI to detect topic changes
        prompt = f"""Analyze these two consecutive paragraphs from a transcript. 
        Determine if there's a significant topic change that warrants a section headline.
        
        Previous paragraph: {prev_para.text[:300]}
        Current paragraph: {current_para.text[:300]}
        
        If there's a major topic change, provide a brief headline (2-4 words) for the current paragraph.
        If not, respond with "NO_HEADLINE".
        
        Response:"""
        
        try:
            response = self.lm_client.generate_text(prompt, max_tokens=50, temperature=0.1)
            if response and response.strip() and response.strip() != "NO_HEADLINE":
                return response.strip()
        except Exception as e:
            print(f"AI topic change detection failed: {e}")
        
        return self._rule_based_topic_change(prev_para, current_para)
    
    def _rule_based_topic_change(self, prev_para: Paragraph, current_para: Paragraph) -> Optional[str]:
        """Rule-based topic change detection"""
        prev_text = prev_para.text.lower()
        current_text = current_para.text.lower()
        
        # Look for topic change indicators
        topic_indicators = [
            'now let\'s talk about', 'speaking of', 'on the topic of',
            'regarding', 'about', 'in terms of', 'when it comes to',
            'moving on to', 'next', 'another', 'different',
            'budget', 'timeline', 'cost', 'planning', 'implementation'
        ]
        
        for indicator in topic_indicators:
            if indicator in current_text:
                # Extract the topic
                if 'budget' in current_text:
                    return "Budget Discussion"
                elif 'timeline' in current_text:
                    return "Timeline Planning"
                elif 'cost' in current_text:
                    return "Cost Analysis"
                elif 'planning' in current_text:
                    return "Project Planning"
                else:
                    return "New Topic"
        
        return None
    
    def _analyze_paragraph_for_headline(self, paragraph: Paragraph, index: int, total: int) -> Optional[str]:
        """Analyze paragraph content to generate appropriate headline"""
        text = paragraph.text.lower()
        
        # Check for explicit topic indicators
        topic_indicators = {
            'introduction': ['introduction', 'intro', 'welcome', 'hello', 'hi'],
            'overview': ['overview', 'summary', 'recap', 'briefly'],
            'background': ['background', 'context', 'history', 'previously'],
            'main topic': ['main', 'primary', 'key', 'important', 'critical'],
            'discussion': ['discussion', 'analysis', 'review', 'examine'],
            'conclusion': ['conclusion', 'summary', 'wrap up', 'finish', 'end'],
            'next steps': ['next', 'following', 'subsequent', 'additional', 'more']
        }
        
        for topic, keywords in topic_indicators.items():
            if any(keyword in text for keyword in keywords):
                return topic.title()
        
        # Check for question patterns
        if '?' in paragraph.text:
            return "Questions & Discussion"
        
        # Check for list patterns
        if any(char in paragraph.text for char in ['•', '-', '*', '1.', '2.', '3.']):
            return "Key Points"
        
        # Check for time-based indicators
        if any(word in text for word in ['first', 'initially', 'start', 'begin']):
            return "Introduction"
        
        if any(word in text for word in ['finally', 'last', 'end', 'conclude']):
            return "Conclusion"
        
        # Default based on position
        if index == 0:
            return "Introduction"
        elif index == total - 1:
            return "Conclusion"
        elif index < total // 3:
            return "Opening Discussion"
        elif index > 2 * total // 3:
            return "Closing Discussion"
        else:
            return "Main Discussion"
    
    def _get_timestamp_marker(self, start_time: str) -> str:
        """Convert SRT timestamp to readable format"""
        # Convert "00:01:23,456" to "1:23"
        time_parts = start_time.split(':')
        if len(time_parts) >= 3:
            hours = int(time_parts[0])
            minutes = int(time_parts[1])
            if hours > 0:
                return f"{hours}:{minutes:02d}"
            else:
                return f"{minutes}:{time_parts[2].split(',')[0]}"
        return start_time


class PythonConverter:
    """Handles conversion to various output formats using Python libraries"""
    
    def __init__(self):
        self.supported_formats = {
            'html': 'html',
            'pdf': 'pdf',
            'docx': 'docx',
            'odt': 'odt',
            'rtf': 'rtf'
        }
        self._check_dependencies()
    
    def _check_dependencies(self):
        """Check which conversion libraries are available"""
        self.available_libs = {
            'weasyprint': False,
            'reportlab': False,
            'markdown': False,
            'docx': False,
            'odfpy': False
        }
        
        # Check for WeasyPrint (PDF from HTML)
        try:
            import weasyprint
            self.available_libs['weasyprint'] = True
        except ImportError:
            pass
        
        # Check for ReportLab (direct PDF)
        try:
            import reportlab
            self.available_libs['reportlab'] = True
        except ImportError:
            pass
        
        # Check for Markdown
        try:
            import markdown
            self.available_libs['markdown'] = True
        except ImportError:
            pass
        
        # Check for python-docx
        try:
            from docx import Document
            self.available_libs['docx'] = True
        except ImportError:
            pass
        
        # Check for odfpy
        try:
            import odf
            self.available_libs['odfpy'] = True
        except ImportError:
            pass
    
    def is_available(self) -> bool:
        """Check if any conversion libraries are available"""
        return any(self.available_libs.values())
    
    def get_available_formats(self) -> List[str]:
        """Get list of available output formats"""
        available = ['markdown']  # Always available
        
        if self.available_libs['markdown']:
            available.append('html')
        
        if self.available_libs['weasyprint'] or self.available_libs['reportlab']:
            available.append('pdf')
        
        if self.available_libs['docx']:
            available.append('docx')
        
        if self.available_libs['odfpy']:
            available.append('odt')
        
        return available
    
    def convert(self, markdown_content: str, output_format: str, output_file: str) -> bool:
        """Convert markdown content to specified format using Python libraries"""
        if output_format not in self.supported_formats:
            log(f"Unsupported format: {output_format}", "ERROR")
            return False
        
        if not self.is_available():
            log("No conversion libraries available. Please install required packages.", "ERROR")
            return False
        
        try:
            if output_format == 'html':
                return self._convert_to_html(markdown_content, output_file)
            elif output_format == 'pdf':
                return self._convert_to_pdf(markdown_content, output_file)
            elif output_format == 'docx':
                return self._convert_to_docx(markdown_content, output_file)
            elif output_format == 'odt':
                return self._convert_to_odt(markdown_content, output_file)
            elif output_format == 'rtf':
                return self._convert_to_rtf(markdown_content, output_file)
            else:
                log(f"Conversion to {output_format} not implemented", "ERROR")
                return False
                
        except Exception as e:
            log(f"Error during conversion: {e}", "ERROR")
            return False
    
    def _convert_to_html(self, markdown_content: str, output_file: str) -> bool:
        """Convert markdown to HTML"""
        if not self.available_libs['markdown']:
            log("Markdown library not available. Install with: pip install markdown", "ERROR")
            return False
        
        try:
            import markdown
            html_content = markdown.markdown(markdown_content, extensions=['tables', 'toc'])
            
            # Wrap in HTML document
            full_html = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <title>Transcript</title>
    <style>
        body {{ font-family: Arial, sans-serif; max-width: 800px; margin: 0 auto; padding: 20px; }}
        h1, h2 {{ color: #333; }}
        h2 {{ border-bottom: 1px solid #eee; padding-bottom: 5px; }}
        a {{ color: #0066cc; text-decoration: none; }}
        a:hover {{ text-decoration: underline; }}
    </style>
</head>
<body>
{html_content}
</body>
</html>"""
            
            with open(output_file, 'w', encoding='utf-8') as f:
                f.write(full_html)
            
            log(f"Successfully converted to HTML")
            return True
        except Exception as e:
            log(f"HTML conversion failed: {e}", "ERROR")
            return False
    
    def _convert_to_pdf(self, markdown_content: str, output_file: str) -> bool:
        """Convert markdown to PDF"""
        # Try WeasyPrint first (better HTML to PDF conversion)
        if self.available_libs['weasyprint']:
            return self._convert_to_pdf_weasyprint(markdown_content, output_file)
        elif self.available_libs['reportlab']:
            return self._convert_to_pdf_reportlab(markdown_content, output_file)
        else:
            log("No PDF conversion libraries available. Install with: pip install weasyprint", "ERROR")
            return False
    
    def _convert_to_pdf_weasyprint(self, markdown_content: str, output_file: str) -> bool:
        """Convert to PDF using WeasyPrint"""
        try:
            import weasyprint
            import markdown
            
            # Convert markdown to HTML first
            html_content = markdown.markdown(markdown_content, extensions=['tables', 'toc'])
            
            # Create HTML document
            html_doc = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <title>Transcript</title>
    <style>
        @page {{ margin: 1in; }}
        body {{ font-family: Arial, sans-serif; line-height: 1.6; }}
        h1, h2 {{ color: #333; page-break-after: avoid; }}
        h2 {{ border-bottom: 1px solid #eee; padding-bottom: 5px; }}
        a {{ color: #0066cc; text-decoration: none; }}
    </style>
</head>
<body>
{html_content}
</body>
</html>"""
            
            # Convert HTML to PDF
            weasyprint.HTML(string=html_doc).write_pdf(output_file)
            log(f"Successfully converted to PDF using WeasyPrint")
            return True
        except Exception as e:
            log(f"WeasyPrint PDF conversion failed: {e}", "ERROR")
            return False
    
    def _convert_to_pdf_reportlab(self, markdown_content: str, output_file: str) -> bool:
        """Convert to PDF using ReportLab"""
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            
            # Create PDF document
            doc = SimpleDocTemplate(output_file, pagesize=letter)
            styles = getSampleStyleSheet()
            story = []
            
            # Custom style for paragraphs
            custom_style = ParagraphStyle(
                'Custom',
                parent=styles['Normal'],
                fontSize=12,
                spaceAfter=12,
            )
            
            # Parse markdown content (simple implementation)
            lines = markdown_content.split('\n')
            for line in lines:
                line = line.strip()
                if not line:
                    story.append(Spacer(1, 6))
                elif line.startswith('# '):
                    story.append(Paragraph(line[2:], styles['Title']))
                elif line.startswith('## '):
                    story.append(Paragraph(line[3:], styles['Heading1']))
                elif line.startswith('**') and line.endswith('**'):
                    story.append(Paragraph(line, styles['Heading2']))
                else:
                    story.append(Paragraph(line, custom_style))
            
            doc.build(story)
            log(f"Successfully converted to PDF using ReportLab")
            return True
        except Exception as e:
            log(f"ReportLab PDF conversion failed: {e}", "ERROR")
            return False
    
    def _convert_to_docx(self, markdown_content: str, output_file: str) -> bool:
        """Convert markdown to DOCX"""
        if not self.available_libs['docx']:
            log("python-docx library not available. Install with: pip install python-docx", "ERROR")
            return False
        
        try:
            from docx import Document
            from docx.shared import Inches
            
            doc = Document()
            
            # Parse markdown content
            lines = markdown_content.split('\n')
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                elif line.startswith('# '):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith('## '):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith('**') and line.endswith('**'):
                    p = doc.add_paragraph()
                    p.add_run(line).bold = True
                else:
                    doc.add_paragraph(line)
            
            doc.save(output_file)
            log(f"Successfully converted to DOCX")
            return True
        except Exception as e:
            log(f"DOCX conversion failed: {e}", "ERROR")
            return False
    
    def _convert_to_odt(self, markdown_content: str, output_file: str) -> bool:
        """Convert markdown to ODT"""
        if not self.available_libs['odfpy']:
            log("odfpy library not available. Install with: pip install odfpy", "ERROR")
            return False
        
        try:
            # ODT conversion is complex, for now just create a simple text document
            # This is a placeholder implementation
            log("ODT conversion not fully implemented. Install odfpy for basic support.", "ERROR")
            return False
        except Exception as e:
            log(f"ODT conversion failed: {e}", "ERROR")
            return False
    
    def _convert_to_rtf(self, markdown_content: str, output_file: str) -> bool:
        """Convert markdown to RTF"""
        try:
            # Simple RTF conversion
            rtf_content = "{\\rtf1\\ansi\\deff0 {\\fonttbl {\\f0 Times New Roman;}}\\f0\\fs24 "
            
            # Basic markdown to RTF conversion
            lines = markdown_content.split('\n')
            for line in lines:
                line = line.strip()
                if not line:
                    rtf_content += "\\par "
                elif line.startswith('# '):
                    rtf_content += f"\\b {line[2:]}\\b0\\par "
                elif line.startswith('## '):
                    rtf_content += f"\\b {line[3:]}\\b0\\par "
                elif line.startswith('**') and line.endswith('**'):
                    rtf_content += f"\\b {line[2:-2]}\\b0\\par "
                else:
                    rtf_content += f"{line}\\par "
            
            rtf_content += "}"
            
            with open(output_file, 'w', encoding='utf-8') as f:
                f.write(rtf_content)
            
            log(f"Successfully converted to RTF")
            return True
        except Exception as e:
            log(f"RTF conversion failed: {e}", "ERROR")
            return False


class MarkdownFormatter:
    """Formats paragraphs and headlines into Markdown"""
    
    def __init__(self, lm_client: Optional[AIClient] = None, video_handler: Optional[VideoURLHandler] = None, no_headlines: bool = False):
        self.headline_generator = HeadlineGenerator(lm_client)
        self.video_handler = video_handler
        self.no_headlines = no_headlines
    
    def format_to_markdown(self, paragraphs: List[Paragraph], title: str = None) -> str:
        """Convert paragraphs to formatted Markdown"""
        if not paragraphs:
            return "# No content found\n"
        
        # Start building markdown
        markdown_parts = []
        
        # Add title
        if title:
            markdown_parts.append(f"# {title}\n")
        else:
            markdown_parts.append("# Transcript\n")
        
        # Generate headlines only if not disabled
        headlines = []
        if not self.no_headlines:
            headlines = self.headline_generator.generate_headlines(paragraphs)
        
        # Add content with or without headlines
        current_headline_index = 0
        
        for i, paragraph in enumerate(paragraphs):
            # Check if we need to add a headline (only if headlines are enabled)
            if (not self.no_headlines and 
                current_headline_index < len(headlines) and 
                headlines[current_headline_index][0] == i):
                _, headline, timestamp = headlines[current_headline_index]
                markdown_parts.append(f"## {headline} *({timestamp})*\n")
                current_headline_index += 1
            
            # Add paragraph with timestamp
            timestamp_marker = self._format_timestamp(paragraph.start_time)
            if self.video_handler and self.video_handler.is_valid():
                # Create clickable timestamp link
                timestamp_url = self.video_handler.get_timestamp_url(paragraph.start_seconds)
                if timestamp_url:
                    timestamp_marker = f"[{timestamp_marker}]({timestamp_url})"
                else:
                    timestamp_marker = f"**{timestamp_marker}**"
            else:
                timestamp_marker = f"**{timestamp_marker}**"
            
            markdown_parts.append(f"{timestamp_marker} {paragraph.text}\n\n")
        
        return ''.join(markdown_parts)
    
    def _slugify(self, text: str) -> str:
        """Convert text to URL-friendly slug"""
        return re.sub(r'[^a-zA-Z0-9\s-]', '', text.lower()).replace(' ', '-')
    
    def _format_timestamp(self, start_time: str) -> str:
        """Format timestamp for display"""
        # Convert "00:01:23,456" to "1:23"
        time_parts = start_time.split(':')
        if len(time_parts) >= 3:
            hours = int(time_parts[0])
            minutes = int(time_parts[1])
            seconds = time_parts[2].split(',')[0]
            if hours > 0:
                return f"{hours}:{minutes:02d}:{seconds}"
            else:
                return f"{minutes}:{seconds}"
        return start_time


class SRTToMarkdownConverter:
    """Main converter class that orchestrates the conversion process"""
    
    def __init__(self, max_gap_seconds: float = 3.0, min_paragraph_length: int = 10, lm_studio_url: str = "http://localhost:1234", model: str = "qwen/qwen3-4b-2507", video_url: str = None, wait_for_model: bool = False, no_headlines: bool = False, output_format: str = "markdown", config_path: str = None, ai_provider: str = None):
        self.parser = SubtitleParser()
        self.config_manager = ConfigManager(config_path)
        self.video_handler = VideoURLHandler(video_url) if video_url else None
        
        # Initialize AI client
        self.ai_client = self._initialize_ai_client(ai_provider, lm_studio_url, model, wait_for_model)
        
        self.grouper = ParagraphGrouper(max_gap_seconds, min_paragraph_length, self.ai_client)
        self.formatter = MarkdownFormatter(self.ai_client, self.video_handler, no_headlines)
        self.python_converter = PythonConverter()
        
        # Use default format from config if not specified
        if output_format == "markdown" and self.config_manager:
            config_output = self.config_manager.get_output_config()
            self.output_format = config_output.get("default_format", output_format)
        else:
            self.output_format = output_format
    
    def _initialize_ai_client(self, ai_provider: str = None, lm_studio_url: str = None, model: str = None, wait_for_model: bool = False) -> Optional[AIClient]:
        """Initialize AI client based on configuration and command line options"""
        # Determine which provider to use
        if ai_provider:
            provider = ai_provider
        else:
            provider = self.config_manager.get_default_provider()
        
        # Get provider configuration
        provider_config = self.config_manager.get_provider_config(provider)
        if not provider_config:
            print(f"⚠ Provider '{provider}' not found in configuration")
            return None
        
        if not provider_config.get("enabled", False):
            print(f"⚠ Provider '{provider}' is disabled in configuration")
            return None
        
        # Override config with command line options if provided
        if lm_studio_url and provider == "lm_studio":
            provider_config = provider_config.copy()
            provider_config["base_url"] = lm_studio_url
        if model:
            provider_config = provider_config.copy()
            provider_config["default_model"] = model
        
        # Check for required API keys
        if provider in ["openai", "anthropic", "grok", "google"]:
            api_key = provider_config.get("api_key", "")
            if not api_key:
                print(f"⚠ API key not configured for {provider}")
                print(f"  Please set the API key in your config file: {self.config_manager.config_path}")
                return None
        
        # Create AI client
        try:
            ai_client = AIClient(provider, provider_config)
            
            # Test availability
            if ai_client.is_available():
                print(f"✓ {provider.title()} is available - using {ai_client.model}")
                if wait_for_model and provider == "lm_studio":
                    if not ai_client.wait_for_model_ready():
                        print("⚠ Model not ready, using fallback processing")
                        return None
                return ai_client
            else:
                print(f"⚠ {provider.title()} is not available")
                return None
                
        except Exception as e:
            print(f"⚠ Error initializing {provider}: {e}")
            return None
    
    def _detect_format_from_extension(self, filename: str) -> str:
        """Detect output format from file extension"""
        if not filename:
            return self.output_format
        
        # Get file extension
        if '.' in filename:
            extension = filename.lower().split('.')[-1]
        else:
            return self.output_format
        
        # Map extensions to formats
        extension_map = {
            'md': 'markdown',
            'markdown': 'markdown',
            'html': 'html',
            'htm': 'html',
            'pdf': 'pdf',
            'docx': 'docx',
            'doc': 'docx',
            'odt': 'odt',
            'rtf': 'rtf'
        }
        
        return extension_map.get(extension, self.output_format)
    
    def convert(self, input_file: str, output_file: str = None, title: str = None) -> str:
        """Convert SRT file to specified format"""
        print(f"Parsing SRT file: {input_file}")
        blocks = self.parser.parse_srt_file(input_file)
        
        if not blocks:
            print("No subtitle blocks found in the file.")
            return ""
        
        print(f"Found {len(blocks)} subtitle blocks")
        
        # Smart format detection from output file extension
        if output_file:
            detected_format = self._detect_format_from_extension(output_file)
            if detected_format != self.output_format:
                print(f"Detected format from filename: {detected_format}")
                self.output_format = detected_format
        
        # Generate default output filename if not provided
        if output_file is None:
            # Remove subtitle extension and add appropriate extension
            if input_file.lower().endswith('.srt'):
                base_name = input_file[:-4]
            elif input_file.lower().endswith('.vtt'):
                base_name = input_file[:-4]
            else:
                base_name = input_file
            
            # Add appropriate extension based on output format
            extensions = {
                'markdown': '.md',
                'html': '.html',
                'pdf': '.pdf',
                'docx': '.docx',
                'odt': '.odt',
                'rtf': '.rtf'
            }
            extension = extensions.get(self.output_format, '.md')
            output_file = base_name + extension
            print(f"Output file: {output_file}")
        
        # Check LM Studio availability
        if self.lm_client and self.lm_client.is_available():
            context_window = self.lm_client.get_context_window()
            print(f"✓ LM Studio is available - using AI-enhanced processing with {self.lm_client.model}")
            print(f"  Model context window: {context_window} tokens")
        else:
            print("⚠ LM Studio not available - using fallback processing")
            if not self.lm_client:
                print("  No suitable model found in LM Studio")
            else:
                print("  Make sure LM Studio is running with a compatible model loaded")
        
        # Check video URL
        if self.video_handler:
            if self.video_handler.is_valid():
                print(f"✓ Video URL configured - timestamps will link to {self.video_handler.platform} video")
            else:
                print("⚠ Invalid video URL - timestamps will not be clickable")
                print(f"  Supported platforms: YouTube, Vimeo, Twitch, Dailymotion, TikTok, Instagram, Facebook, Twitter, LinkedIn, Rumble, Odysee")
                print(f"  Provided URL: {self.video_handler.video_url}")
        
        print("Grouping into paragraphs...")
        try:
            paragraphs = self.grouper.group_into_paragraphs(blocks)
            print(f"Created {len(paragraphs)} paragraphs")
            
            if not paragraphs:
                print("⚠ No paragraphs created from subtitle blocks")
                return ""
        except Exception as e:
            print(f"Error during paragraph grouping: {e}")
            return ""
        
        print("Generating headlines...")
        try:
            markdown_content = self.formatter.format_to_markdown(paragraphs, title)
            
            if not markdown_content.strip():
                print("⚠ Generated markdown content is empty")
                return ""
        except Exception as e:
            print(f"Error during markdown formatting: {e}")
            return ""
        
        try:
            if self.output_format == 'markdown':
                # Direct markdown output
                with open(output_file, 'w', encoding='utf-8') as f:
                    f.write(markdown_content)
                print(f"Markdown saved to: {output_file}")
            else:
                # Convert using Python libraries
                if self.python_converter.is_available():
                    if self.python_converter.convert(markdown_content, self.output_format, output_file):
                        print(f"{self.output_format.upper()} saved to: {output_file}")
                    else:
                        print(f"Failed to convert to {self.output_format.upper()}")
                        return ""
                else:
                    print(f"Python conversion libraries not available. Saving as markdown instead.")
                    # Fallback to markdown
                    markdown_file = output_file.rsplit('.', 1)[0] + '.md'
                    with open(markdown_file, 'w', encoding='utf-8') as f:
                        f.write(markdown_content)
                    print(f"Markdown saved to: {markdown_file}")
                    return markdown_content
        except Exception as e:
            print(f"Error writing output file: {e}")
            return ""
        
        return markdown_content


def setup_configuration(config_path: str = None) -> None:
    """Interactive configuration setup"""
    config_manager = ConfigManager(config_path)
    
    print("🔧 Sub2Trans Configuration Setup")
    print("=" * 40)
    print(f"Config file: {config_manager.config_path}")
    print()
    
    # Get default provider
    current_default = config_manager.get_default_provider()
    print(f"Current default provider: {current_default}")
    
    providers = ["lm_studio", "openai", "anthropic", "grok", "google"]
    print("\nAvailable providers:")
    for i, provider in enumerate(providers, 1):
        config = config_manager.get_provider_config(provider)
        status = "✓ Enabled" if config.get("enabled", False) else "✗ Disabled"
        print(f"  {i}. {provider.title()} - {status}")
    
    try:
        choice = input(f"\nSelect default provider (1-{len(providers)}, or press Enter to keep '{current_default}'): ").strip()
        if choice:
            provider_idx = int(choice) - 1
            if 0 <= provider_idx < len(providers):
                config_manager.config["default_provider"] = providers[provider_idx]
                print(f"✓ Set default provider to {providers[provider_idx]}")
        
        # Configure API keys for cloud providers
        for provider in ["openai", "anthropic", "grok", "google"]:
            config = config_manager.get_provider_config(provider)
            if not config.get("api_key"):
                print(f"\n{provider.title()} API Key:")
                print("  Leave empty to skip, or enter your API key")
                api_key = input(f"  API Key: ").strip()
                if api_key:
                    config_manager.update_provider_config(provider, {"api_key": api_key, "enabled": True})
                    print(f"✓ {provider.title()} configured")
                else:
                    print(f"  Skipped {provider.title()}")
        
        # Configure default output format
        print(f"\nOutput Format Configuration:")
        current_format = config_manager.get_output_config().get("default_format", "markdown")
        print(f"Current default format: {current_format}")
        
        formats = ["markdown", "html", "pdf", "docx", "odt", "rtf"]
        print("\nAvailable formats:")
        for i, fmt in enumerate(formats, 1):
            print(f"  {i}. {fmt.upper()}")
        
        try:
            format_choice = input(f"\nSelect default format (1-{len(formats)}, or press Enter to keep '{current_format}'): ").strip()
            if format_choice:
                format_idx = int(format_choice) - 1
                if 0 <= format_idx < len(formats):
                    config_manager.config["output"]["default_format"] = formats[format_idx]
                    print(f"✓ Set default format to {formats[format_idx]}")
        except (ValueError, IndexError):
            print("  Invalid choice, keeping current format")
        
        # Save configuration
        config_manager._save_config()
        print(f"\n✓ Configuration saved to {config_manager.config_path}")
        
    except KeyboardInterrupt:
        print("\n\nSetup cancelled")
    except Exception as e:
        print(f"\nError during setup: {e}")


def show_configuration(config_path: str = None) -> None:
    """Display current configuration"""
    config_manager = ConfigManager(config_path)
    
    print("📋 Sub2Trans Current Configuration")
    print("=" * 40)
    print(f"Config file: {config_manager.config_path}")
    print()
    
    # Show default provider
    default_provider = config_manager.get_default_provider()
    print(f"Default AI Provider: {default_provider.title()}")
    
    # Show provider status
    print("\nAI Provider Status:")
    for provider, config in config_manager.config.get("ai_providers", {}).items():
        status = "✓ Enabled" if config.get("enabled", False) else "✗ Disabled"
        has_key = "✓ API Key Set" if config.get("api_key") else "✗ No API Key"
        print(f"  {provider.title()}: {status} ({has_key})")
    
    # Show output configuration
    output_config = config_manager.get_output_config()
    print(f"\nOutput Configuration:")
    print(f"  Default Format: {output_config.get('default_format', 'markdown')}")
    print(f"  Pandoc Path: {output_config.get('pandoc_path', 'pandoc')}")
    print(f"  Auto-detect Format: {output_config.get('auto_detect_format', True)}")
    
    # Show processing configuration
    processing_config = config_manager.get_processing_config()
    print(f"\nProcessing Configuration:")
    print(f"  Max Gap Seconds: {processing_config.get('max_gap_seconds', 3.0)}")
    print(f"  Min Paragraph Length: {processing_config.get('min_paragraph_length', 10)}")
    print(f"  Wait for Model: {processing_config.get('wait_for_model', False)}")
    print(f"  No Headlines: {processing_config.get('no_headlines', False)}")
    print(f"  Verbose: {processing_config.get('verbose', False)}")


def list_available_formats() -> None:
    """Display available output formats and required libraries"""
    converter = PythonConverter()
    
    print("📄 Available Output Formats")
    print("=" * 40)
    
    # Always available formats
    print("✅ Always Available:")
    print("  • markdown - Native format, no dependencies")
    
    # Check each format
    if converter.available_libs['markdown']:
        print("\n✅ HTML Conversion:")
        print("  • html - Requires: pip install markdown")
    else:
        print("\n❌ HTML Conversion:")
        print("  • html - Install with: pip install markdown")
    
    if converter.available_libs['weasyprint']:
        print("\n✅ PDF Conversion (WeasyPrint):")
        print("  • pdf - Requires: pip install weasyprint")
    elif converter.available_libs['reportlab']:
        print("\n✅ PDF Conversion (ReportLab):")
        print("  • pdf - Requires: pip install reportlab")
    else:
        print("\n❌ PDF Conversion:")
        print("  • pdf - Install with: pip install weasyprint (recommended) or pip install reportlab")
    
    if converter.available_libs['docx']:
        print("\n✅ DOCX Conversion:")
        print("  • docx - Requires: pip install python-docx")
    else:
        print("\n❌ DOCX Conversion:")
        print("  • docx - Install with: pip install python-docx")
    
    if converter.available_libs['odfpy']:
        print("\n✅ ODT Conversion:")
        print("  • odt - Requires: pip install odfpy")
    else:
        print("\n❌ ODT Conversion:")
        print("  • odt - Install with: pip install odfpy")
    
    print("\n✅ RTF Conversion:")
    print("  • rtf - Native support, no dependencies")
    
    print(f"\n📋 Summary:")
    available_formats = converter.get_available_formats()
    print(f"  Available formats: {', '.join(available_formats)}")
    
    if len(available_formats) == 1:
        print("  💡 Install additional libraries to enable more formats!")
    else:
        print(f"  🎉 {len(available_formats)} formats available!")


def main():
    """Command-line interface for the converter"""
    parser = argparse.ArgumentParser(
        description="Convert subtitle files (SRT, VTT, ASS, SSA, SBV, TTML, SAMI, LRC) to well-formatted transcripts",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  python sub2trans.py input.srt                    # Creates input.md
  python sub2trans.py input.vtt                    # Creates input.md (VTT support)
  python sub2trans.py input.ass                  # Creates input.md (ASS support)
  python sub2trans.py input.sbv                  # Creates input.md (SBV support)
  python sub2trans.py input.ttml                 # Creates input.md (TTML support)
  python sub2trans.py input.srt -o output.md      # Creates output.md
  python sub2trans.py input.vtt -o output.pdf      # Smart format detection (PDF)
  python sub2trans.py input.srt -t "My Video Transcript"  # Creates input.md with title
  python sub2trans.py input.srt --max-gap 5.0 --min-length 30
  python sub2trans.py input.srt --lm-studio-url http://localhost:1234 --model qwen/qwen3-4b-2507
  python sub2trans.py input.srt --no-ai
  python sub2trans.py input.srt -u "https://www.youtube.com/watch?v=dQw4w9WgXcQ"
  python sub2trans.py input.srt -u "https://vimeo.com/123456789" -o transcript.md
  python sub2trans.py --list-models                    # List available models
  python sub2trans.py input.srt --model "llama-2-7b"  # Use specific model
  python sub2trans.py input.srt --wait-for-model      # Wait for model to load
  python sub2trans.py input.srt -f pdf                # Convert to PDF
  python sub2trans.py input.srt -f docx               # Convert to Word document
  python sub2trans.py input.srt -f html               # Convert to HTML
  python sub2trans.py input.srt -o output.pdf         # Smart format detection (PDF)
  python sub2trans.py input.srt -o output.docx        # Smart format detection (Word)
  python sub2trans.py input.srt -o output.html        # Smart format detection (HTML)

AI Integration:
  This tool supports multiple AI providers for enhanced processing:
  - LM Studio (local): Free, runs locally with your own models
  - OpenAI (ChatGPT): Cloud-based, requires API key
  - Anthropic (Claude): Cloud-based, requires API key  
  - Grok (X.AI): Cloud-based, requires API key
  - Google (Gemini): Cloud-based, requires API key
  
  Features:
  - Better paragraph grouping for long transcripts
  - AI-generated headlines and section titles
  - Improved content organization
  - Automatic model selection and fallback
  
  Configuration:
  - Use --setup-config to configure AI providers and default output format
  - Config file: ~/sub2trans_config.json
  - Use --ai-provider to override the default provider
  - Set default output format in config file (markdown, html, pdf, docx, odt, rtf)

Video URL Integration:
  Use -u/--video-url to create clickable timestamps that link to the video:
  - YouTube: Creates links with ?t=120s format
  - Vimeo: Creates links with #t=120.521 format
  - Twitch: Creates links with ?t=120s format
  - Dailymotion: Creates links with ?start=120 format
  - Facebook: Creates links with ?t=120 format
  - Rumble: Creates links with ?t=120 format
  - Odysee: Creates links with ?t=120 format
  - TikTok, Instagram, Twitter, LinkedIn: Base URLs (no timestamp support)
  - Timestamps will be clickable in the generated Markdown
        """
    )
    
    parser.add_argument('input_file', nargs='?', help='Input SRT or VTT file path')
    parser.add_argument('-o', '--output', help='Output file path (default: input file with .md extension). Smart format detection: .pdf, .docx, .html, .odt, .rtf')
    parser.add_argument('-t', '--title', help='Title for the document')
    parser.add_argument('--max-gap', type=float, default=3.0,
                       help='Maximum gap in seconds between subtitles to group into paragraphs (default: 3.0)')
    parser.add_argument('--min-length', type=int, default=10,
                       help='Minimum paragraph length in characters (default: 10)')
    parser.add_argument('--preview', action='store_true',
                       help='Preview the first few paragraphs without saving')
    parser.add_argument('--lm-studio-url', default='http://localhost:1234',
                       help='LM Studio API URL (default: http://localhost:1234)')
    parser.add_argument('--model', default='qwen/qwen3-4b-2507',
                       help='Model name to use with LM Studio (default: qwen/qwen3-4b-2507)')
    parser.add_argument('--wait-for-model', action='store_true',
                       help='Wait for model to load before processing (useful when switching models)')
    parser.add_argument('--no-ai', action='store_true',
                       help='Disable AI processing and use fallback methods only')
    parser.add_argument('-v', '--verbose', action='store_true',
                       help='Enable verbose output')
    parser.add_argument('-q', '--quiet', action='store_true',
                       help='Suppress non-essential output')
    parser.add_argument('--batch', action='store_true',
                       help='Process all SRT files in current directory')
    parser.add_argument('--no-headlines', action='store_true',
                       help='Disable headline generation (output raw transcript)')
    parser.add_argument('-f', '--format', choices=['markdown', 'html', 'pdf', 'docx', 'odt', 'rtf'], 
                       default='markdown', help='Output format (default: markdown)')
    parser.add_argument('--list-formats', action='store_true',
                       help='List available output formats and required libraries')
    parser.add_argument('-u', '--video-url', 
                       help='Video URL (YouTube or Vimeo) to create clickable timestamps')
    parser.add_argument('--list-models', action='store_true',
                       help='List available models in LM Studio and exit')
    parser.add_argument('--config', help='Path to configuration file (default: ~/sub2trans_config.json)')
    parser.add_argument('--ai-provider', choices=['lm_studio', 'openai', 'anthropic', 'grok', 'google'],
                       help='AI provider to use (overrides config file)')
    parser.add_argument('--setup-config', action='store_true',
                       help='Setup configuration file interactively')
    parser.add_argument('--show-config', action='store_true',
                       help='Show current configuration and exit')
    
    args = parser.parse_args()
    
    # Set global verbosity
    global VERBOSE
    VERBOSE = args.verbose
    
    # Handle setup-config option
    if args.setup_config:
        setup_configuration(args.config)
        sys.exit(0)
    
    # Handle show-config option
    if args.show_config:
        show_configuration(args.config)
        sys.exit(0)
    
    # Handle list-formats option
    if args.list_formats:
        list_available_formats()
        sys.exit(0)
    
    # Handle batch processing
    if args.batch:
        import glob
        subtitle_files = glob.glob("*.srt") + glob.glob("*.vtt")
        if not subtitle_files:
            print("No SRT or VTT files found in current directory")
            sys.exit(1)

        print(f"Found {len(subtitle_files)} subtitle files to process")
        for subtitle_file in subtitle_files:
            print(f"\nProcessing {subtitle_file}...")
            try:
                converter = SRTToMarkdownConverter(args.max_gap, args.min_length, args.lm_studio_url, args.model, args.video_url, args.wait_for_model, args.no_headlines, args.format, args.pandoc_path)
                converter.convert(subtitle_file, title=args.title)
                print(f"✓ {subtitle_file} processed successfully")
            except Exception as e:
                print(f"✗ Error processing {subtitle_file}: {e}")
        sys.exit(0)
    
    # Validate arguments
    if not args.list_models and not args.input_file:
        parser.error("input_file is required unless using --list-models")
    
    # Handle --list-models option
    if args.list_models:
        # This section needs to be updated to use the new AI client system
        # For now, we'll skip the list-models functionality for non-LM Studio providers
        if args.ai_provider and args.ai_provider != "lm_studio":
            print(f"List models not supported for {args.ai_provider}")
            sys.exit(0)
        
        # Create a temporary AI client for LM Studio
        config_manager = ConfigManager(args.config)
        provider_config = config_manager.get_provider_config("lm_studio")
        if not provider_config:
            print("LM Studio not configured")
            sys.exit(1)
        
        # Override with command line options
        if args.lm_studio_url:
            provider_config = provider_config.copy()
            provider_config["base_url"] = args.lm_studio_url
        if args.model:
            provider_config = provider_config.copy()
            provider_config["default_model"] = args.model
        
        lm_client = AIClient("lm_studio", provider_config)
        if lm_client.is_available():
            models = lm_client.get_available_models()
            if models:
                print("Available models in LM Studio:")
                for i, model in enumerate(models, 1):
                    print(f"  {i}. {model}")
            else:
                print("No models found in LM Studio")
        else:
            print("LM Studio is not available")
            print(f"Make sure LM Studio is running at {args.lm_studio_url}")
        sys.exit(0)
    
    # Determine output format (use config default if not specified)
    output_format = args.format
    if output_format == "markdown":  # Only use config default if user didn't specify a format
        try:
            config_manager = ConfigManager(args.config)
            config_output = config_manager.get_output_config()
            output_format = config_output.get("default_format", "markdown")
            if VERBOSE:
                print(f"Using default format from config: {output_format}")
        except:
            pass  # Keep original format if config fails
    
    # Create converter
    if args.no_ai:
        # Create a dummy AI client that's not available
        class DummyAIClient:
            def is_available(self): return False
        dummy_client = DummyAIClient()
        converter = SRTToMarkdownConverter(args.max_gap, args.min_length, "dummy", "dummy", args.video_url, False, args.no_headlines, output_format, args.config, None)
        converter.ai_client = dummy_client
        converter.grouper.lm_client = dummy_client
        converter.formatter.headline_generator.lm_client = dummy_client
    else:
        converter = SRTToMarkdownConverter(args.max_gap, args.min_length, args.lm_studio_url, args.model, args.video_url, args.wait_for_model, args.no_headlines, output_format, args.config, args.ai_provider)
    
    try:
        if args.preview:
            # Preview mode - show first few paragraphs
            try:
                blocks = converter.parser.parse_srt_file(args.input_file)
                if not blocks:
                    print("No subtitle blocks found in the file")
                    sys.exit(1)
                
                paragraphs = converter.grouper.group_into_paragraphs(blocks)
                if not paragraphs:
                    print("No paragraphs created from subtitle blocks")
                    sys.exit(1)
                
                print(f"\nPreview of first 3 paragraphs:\n")
                for i, para in enumerate(paragraphs[:3]):
                    timestamp = converter.formatter._format_timestamp(para.start_time)
                    print(f"**{timestamp}** {para.text}\n")
                
                if len(paragraphs) > 3:
                    print(f"... and {len(paragraphs) - 3} more paragraphs")
            except Exception as e:
                print(f"Error in preview mode: {e}")
                if VERBOSE:
                    import traceback
                    traceback.print_exc()
                sys.exit(1)
        else:
            # Full conversion
            try:
                result = converter.convert(args.input_file, args.output, args.title)
                if not result:
                    print("Conversion failed - no output generated")
                    sys.exit(1)
            except Exception as e:
                print(f"Error during conversion: {e}")
                if VERBOSE:
                    import traceback
                    traceback.print_exc()
                sys.exit(1)
    
    except KeyboardInterrupt:
        print("\nOperation cancelled by user")
        sys.exit(1)
    except FileNotFoundError as e:
        print(f"File not found: {e}")
        sys.exit(1)
    except PermissionError as e:
        print(f"Permission denied: {e}")
        sys.exit(1)
    except Exception as e:
        print(f"Unexpected error: {e}")
        if VERBOSE:
            import traceback
            traceback.print_exc()
        sys.exit(1)


if __name__ == "__main__":
    main()

